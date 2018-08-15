from docx import Document
import re
import itertools
import parsepdf

# https://python-docx.readthedocs.io/en/latest/


def extract_docx(filepath):
    document = Document(filepath)
    paragraph_list = [
        strip_non_breaking_space(i.text) for i in document.paragraphs
    ]

    table_list = []

    def text_from_table(table):
        for row in table.rows:
            for cell in row.cells:
                table_list.append(strip_non_breaking_space(cell.text))

    for i in document.tables:
        text_from_table(i)

    return paragraph_list + table_list


def extract_pdf(filepath):
    return [parsepdf.pdf_to_text(filepath).replace('\n', ' ')]


def strip_non_breaking_space(sentence):
    return sentence.replace('\xa0', ' ')


def start_extract(filepath, stared=False):
    file_type = filepath.split('.')[-1].lower()
    if 'docx' in file_type:
        full_text = extract_docx(filepath)
    elif 'pdf' in file_type:
        full_text = extract_pdf(filepath)

    def check_lawnet_compatibility(case_citation):
        lawnet_casetypes = ['SLR', 'SGCA', 'SGHC', 'WLR', 'MLJ', 'Ch', 'SGDC', 'AC', ' UKSC', 'Ch', 'WLR', 'QB']
        for casetype in lawnet_casetypes:
            return True

    if stared:
        citation_pattern = re.compile(r'\*[^\[\]]*([[(][1-2]\d{3}(?:-[1-2]\d{3})?[\])]\s[\d\s]*[LR]+\s\d+\s[EqCP]+\s+\d+)|\*[^\[\]]*([[(][1-2]\d{3}(?:-[1-2]\d{3})?[\])]\s[\d\s]*[SLR()WMJChAFQBtra]+\s\d+)|\*[^\[\]]*(\[[1-2]\d{3}(?:-[1-2]\d{3})?\]\s[A-Za-z()]+\s\d+)'
        )
    else:
        citation_pattern = re.compile(r'[\[\(][1-2]\d{3}(?:\-[1-2]\d{3})?[\]\)]\s[\d\s]*[LR]+\s\d+\s[EqCP]+\s+\d+|[\[\(][1-2]\d{3}(?:\-[1-2]\d{3})?[\]\)]\s[\d\s]*[SLR()WLRMLJChACFQBStra]+\s\d+|\[[1-2]\d{3}(?:\-[1-2]\d{3})?\]\s[A-Za-z()]+\s\d+'
        )

    citation_list = [re.findall(citation_pattern, i) for i in full_text]
    if stared:
        citation_list = [
            item[1] for item in itertools.chain.from_iterable(citation_list)
        ]
    else:
        citation_list = itertools.chain.from_iterable(citation_list)

    citation_list = set(citation_list)

    citation_list = [
        ' '.join(citation.split()) for citation in citation_list
    ]
    return citation_list
